"""Read a Word document and output text content.
Supports: .docx, .doc, .rtf, .odt

Improvements over read-docx.py:
  - Extracts headers and footers from all document sections
  - Preserves heading hierarchy (# / ## markers)
  - Better .doc fallback chain: antiword -> docx2txt -> LibreOffice -> raw
  - Handles both linked and unlinked headers/footers correctly
  - raw .doc reading is last resort only (produces garbage for modern .doc)

Usage: python read-docx.py <file_path>
"""
import sys
import os
import subprocess

if len(sys.argv) < 2:
    print("Usage: python read-docx.py <file_path>")
    sys.exit(1)

path = sys.argv[1]
ext = os.path.splitext(path)[1].lower()


def extract_header_footer_texts(doc):
    """Extract text from all section headers and footers."""
    texts = []
    hf_attrs = (
        'header', 'footer',
        'even_page_header', 'even_page_footer',
        'first_page_header', 'first_page_footer',
    )
    seen = set()
    for section in doc.sections:
        for attr in hf_attrs:
            try:
                hf = getattr(section, attr)
                if hf is None or hf.is_linked_to_previous:
                    continue
                for para in hf.paragraphs:
                    t = para.text.strip()
                    if t and t not in seen:
                        texts.append(t)
                        seen.add(t)
            except Exception:
                pass
    return texts


if ext == '.docx':
    try:
        from docx import Document
        doc = Document(path)
        output = []

        # Headers and footers -- often contain document title, date, company name
        hf = extract_header_footer_texts(doc)
        if hf:
            output.append("[Document Header/Footer]")
            output.extend(hf)
            output.append("")

        # Main body paragraphs with heading markers
        for para in doc.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading "):
                level_str = style_name.replace("Heading ", "").strip()
                level = int(level_str) if level_str.isdigit() else 1
                output.append(f"{'#' * level} {t}")
            else:
                output.append(t)

        # Tables
        for i, table in enumerate(doc.tables):
            output.append(f"\n[Table {i + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                output.append("\t".join(cells))

        print("\n".join(output) if output else "[No text content found]")
        sys.exit(0)
    except ImportError:
        print("[ERROR] Cannot read DOCX: install python-docx\n  pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] Failed to read DOCX: {e}")
        sys.exit(1)

elif ext == '.doc':
    # Strategy 1: antiword (best for legacy .doc)
    try:
        result = subprocess.run(['antiword', path], capture_output=True, text=True, timeout=30)
        if result.returncode == 0 and result.stdout.strip():
            print(result.stdout)
            sys.exit(0)
    except Exception:
        pass

    # Strategy 2: docx2txt (pure Python, handles most .doc files)
    try:
        import docx2txt
        text = docx2txt.process(path)
        if text and text.strip():
            print(text)
            sys.exit(0)
    except ImportError:
        pass

    # Strategy 3: LibreOffice -- convert to .docx then re-read properly
    try:
        import tempfile
        tmp_dir = tempfile.mkdtemp()
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', tmp_dir, path],
            capture_output=True, text=True, timeout=90
        )
        if result.returncode == 0:
            converted = os.path.join(tmp_dir, os.path.splitext(os.path.basename(path))[0] + '.docx')
            if os.path.exists(converted):
                result2 = subprocess.run(
                    [sys.executable, __file__, converted],
                    capture_output=True, text=True
                )
                print(result2.stdout)
                sys.exit(result2.returncode)
    except Exception:
        pass

    print("[ERROR] Cannot read .doc file. Install antiword, docx2txt, or LibreOffice.")
    sys.exit(1)

elif ext == '.rtf':
    try:
        from striprtf.striprtf import rtf_to_text
        with open(path, 'r', errors='ignore') as f:
            print(rtf_to_text(f.read()))
        sys.exit(0)
    except ImportError:
        with open(path, 'r', errors='ignore') as f:
            print(f.read())
        sys.exit(0)

else:
    # .odt and other text-like formats
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        print(f.read())
    sys.exit(0)
